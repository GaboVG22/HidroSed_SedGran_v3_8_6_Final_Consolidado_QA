from __future__ import annotations

import io
import numpy as np
import pandas as pd


def _sf(x, default=np.nan):
    try:
        v = float(x)
        return v if np.isfinite(v) else default
    except Exception:
        return default


def audit_external_report(data: dict | pd.DataFrame) -> pd.DataFrame:
    """Generic external-report audit. Accepts dict flags or an Excel-like dataframe.

    Expected fields are intentionally broad so the tool can be used for any project.
    """
    rows = []
    if isinstance(data, pd.DataFrame):
        d = {str(r.get("campo", r.get("item", ""))).strip(): r.get("valor", r.get("estado", "")) for _, r in data.iterrows()}
    else:
        d = data or {}

    def present(key):
        v = d.get(key, d.get(key.lower(), None))
        if v is None:
            return False
        if isinstance(v, str):
            return v.strip().lower() not in ["", "no", "false", "0", "sin dato", "n/a"]
        return bool(v)

    checks = [
        ("datos_cuenca", "Datos base y morfometría", "alta", "Completar área, longitud, cotas, pendiente y régimen."),
        ("idf", "Precipitación e IDF", "alta", "Incorporar P24(T), coeficientes de duración y estación representativa."),
        ("tiempo_concentracion", "Tiempo de concentración", "media", "Comparar métodos y justificar método rector."),
        ("caudales_metodos", "Caudales por métodos", "alta", "Calcular más de un método o justificar método único."),
        ("caudal_adoptado", "Selección del caudal de diseño", "crítica", "Indicar criterio y verificar que valor adoptado coincida."),
        ("hidraulica", "Modelación hidráulica", "alta", "Modelar secciones, WSE, velocidades, Froude y desbordes."),
        ("condicion_aguas_abajo", "Condición aguas abajo", "media", "Revisar marea, humedal, laguna, barra o cota fija si aplica."),
        ("sedimentos_socavacion", "Sedimentos y socavación", "alta", "Incluir granulometría, transporte y socavación general/local."),
        ("proteccion_fluvial", "Protección fluvial", "media", "Predimensionar y justificar enrocado/gaviones/filtro/pie."),
        ("trazabilidad", "QA y trazabilidad", "alta", "Respaldar fuentes, unidades, fórmulas y parámetros críticos."),
    ]
    for key, item, sev, rec in checks:
        ok = present(key)
        rows.append({
            "item": item,
            "campo": key,
            "estado": "OK" if ok else "OBSERVADO",
            "severidad": "baja" if ok else sev,
            "observacion": "Información presente" if ok else f"Falta o no queda trazable: {item}",
            "recomendacion": "Sin acción" if ok else rec,
        })

    # Specific consistency flags.
    if present("criterio_promedio") and present("adopta_maximo"):
        rows.append({"item": "Caudal adoptado", "campo": "contradiccion_promedio_maximo", "estado": "OBSERVADO", "severidad": "crítica", "observacion": "Se declara promedio pero se adopta máximo/envolvente.", "recomendacion": "Corregir criterio o justificar adopción conservadora."})
    if present("p24_constante_verni_king"):
        rows.append({"item": "Verni-King", "campo": "p24_constante_verni_king", "estado": "OBSERVADO", "severidad": "alta", "observacion": "Se usa P24 constante donde la fórmula exige P24(T).", "recomendacion": "Recalcular con P24 por período o marcar como auditoría documental."})
    if present("errores_unidades"):
        rows.append({"item": "Unidades", "campo": "errores_unidades", "estado": "OBSERVADO", "severidad": "crítica", "observacion": "Se detectan inconsistencias de unidades.", "recomendacion": "Normalizar unidades SI y recalcular tablas afectadas."})
    if present("puente_sin_socavacion_local"):
        rows.append({"item": "Socavación local", "campo": "puente_sin_socavacion_local", "estado": "OBSERVADO", "severidad": "crítica", "observacion": "Proyecto de puente sin análisis de pilas/estribos/contracción.", "recomendacion": "Incorporar socavación local y verificar fundaciones."})
    return pd.DataFrame(rows)


DEFAULT_WEIGHTS = {
    "Datos base y morfometría": 0.10,
    "Precipitación e IDF": 0.10,
    "Tiempo de concentración": 0.10,
    "Caudales por métodos": 0.15,
    "Selección del caudal de diseño": 0.15,
    "Modelación hidráulica": 0.15,
    "Condición aguas abajo": 0.05,
    "Sedimentos y socavación": 0.10,
    "Protección fluvial": 0.05,
    "QA y trazabilidad": 0.05,
}


def technical_score(audit_df: pd.DataFrame, module_scores: dict | None = None, threshold: float = 8.7) -> pd.DataFrame:
    """Scores 0-10 with blocking caps from critical rules."""
    module_scores = module_scores or {}
    rows = []
    for module, w in DEFAULT_WEIGHTS.items():
        base = _sf(module_scores.get(module, 10.0), 10.0)
        if audit_df is not None and not audit_df.empty:
            obs = audit_df[(audit_df["item"].astype(str) == module) & (audit_df["estado"].astype(str) != "OK")]
            if not obs.empty:
                maxsev = obs["severidad"].astype(str).str.lower().tolist()
                if "crítica" in maxsev or "critica" in maxsev:
                    base = min(base, 5.0)
                elif "alta" in maxsev:
                    base = min(base, 6.5)
                elif "media" in maxsev:
                    base = min(base, 8.0)
                else:
                    base = min(base, 9.0)
        rows.append({"modulo": module, "ponderacion": w, "nota_modulo": base, "aporte": base * w})
    detail = pd.DataFrame(rows)
    score = float(detail["aporte"].sum())
    cap = 10.0
    cap_reasons = []
    if audit_df is not None and not audit_df.empty:
        fields = ";".join(audit_df.get("campo", pd.Series(dtype=str)).astype(str).tolist()).lower()
        if "contradiccion_promedio_maximo" in fields:
            cap = min(cap, 8.4); cap_reasons.append("contradicción criterio declarado/adoptado")
        if "p24_constante_verni_king" in fields:
            cap = min(cap, 8.0); cap_reasons.append("error metodológico en caudales")
        if "hidraulica" in fields and (audit_df[(audit_df.get("campo", "").astype(str)=="hidraulica") & (audit_df.get("estado", "").astype(str)!="OK")].shape[0] > 0):
            cap = min(cap, 7.8); cap_reasons.append("falta modelación hidráulica suficiente")
        if "puente_sin_socavacion_local" in fields:
            cap = min(cap, 7.5); cap_reasons.append("falta análisis de socavación local en puente")
        if "errores_unidades" in fields:
            cap = min(cap, 7.8); cap_reasons.append("errores de unidades")
        if "trazabilidad" in fields and (audit_df[(audit_df.get("campo", "").astype(str)=="trazabilidad") & (audit_df.get("estado", "").astype(str)!="OK")].shape[0] > 0):
            cap = min(cap, 8.0); cap_reasons.append("falta trazabilidad")
    score_capped = min(score, cap)
    if score_capped >= threshold:
        estado = "Aprobado técnicamente"
    elif score_capped >= 8.2:
        estado = "Aprobado con observaciones menores"
    elif score_capped >= 7.5:
        estado = "Aprobado con observaciones relevantes"
    elif score_capped >= 6.0:
        estado = "Requiere corrección"
    else:
        estado = "No recomendado técnicamente"
    summary = pd.DataFrame([{
        "nota_global": round(score_capped, 2),
        "nota_sin_bloqueos": round(score, 2),
        "tope_aplicado": cap,
        "razones_tope": "; ".join(cap_reasons) or "sin topes",
        "umbral_usuario": threshold,
        "supera_umbral": bool(score_capped >= threshold),
        "estado_final": estado,
    }])
    return pd.concat([summary.assign(tipo="resumen"), detail.assign(tipo="detalle")], ignore_index=True, sort=False)


def excel_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        for name, df in sheets.items():
            if df is not None and isinstance(df, pd.DataFrame) and not df.empty:
                df.to_excel(writer, sheet_name=str(name)[:31], index=False)
    return bio.getvalue()


def technical_markdown_report(title: str, audit_df: pd.DataFrame, score_df: pd.DataFrame, adoption_df: pd.DataFrame | None = None) -> str:
    lines = [f"# {title}", "", "## Resumen técnico"]
    if score_df is not None and not score_df.empty and "nota_global" in score_df.columns:
        row = score_df[score_df.get("tipo", "") == "resumen"].head(1)
        if row.empty:
            row = score_df.head(1)
        r = row.iloc[0]
        lines += [f"- Nota global: **{r.get('nota_global', 'NA')} / 10**", f"- Estado final: **{r.get('estado_final', 'NA')}**", f"- Supera umbral: **{r.get('supera_umbral', 'NA')}**", ""]
    if audit_df is not None and not audit_df.empty:
        lines += ["## Observaciones de auditoría", ""]
        for _, r in audit_df.iterrows():
            if str(r.get("estado")) != "OK":
                lines.append(f"- **{r.get('severidad','')}** · {r.get('item','')}: {r.get('observacion','')}. Recomendación: {r.get('recomendacion','')}")
        lines.append("")
    if adoption_df is not None and not adoption_df.empty:
        lines += ["## Caudal adoptado", "", "```csv", adoption_df.to_csv(index=False).strip(), "```", ""]
    lines += ["## Nota", "Informe generado automáticamente por HidroSed SedGran v3.7.6. Los resultados son preliminares y deben contrastarse con topografía, normativa y revisión profesional."]
    return "\n".join(lines)


def docx_report_bytes(title: str, audit_df: pd.DataFrame, score_df: pd.DataFrame, adoption_df: pd.DataFrame | None = None) -> bytes:
    """Creates a Word report if python-docx is installed."""
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph("Informe automático HidroSed SedGran v3.7.6. Resultados preliminares sujetos a revisión profesional y normativa aplicable.")
    if score_df is not None and not score_df.empty:
        doc.add_heading("Nota técnica", level=1)
        row = score_df[score_df.get("tipo", "") == "resumen"].head(1)
        if row.empty: row = score_df.head(1)
        r = row.iloc[0]
        doc.add_paragraph(f"Nota global: {r.get('nota_global', 'NA')} / 10")
        doc.add_paragraph(f"Estado final: {r.get('estado_final', 'NA')}")
        doc.add_paragraph(f"Razones de tope: {r.get('razones_tope', 'NA')}")
    if audit_df is not None and not audit_df.empty:
        doc.add_heading("Observaciones de auditoría", level=1)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        for i, h in enumerate(["Item", "Estado", "Severidad", "Observación", "Recomendación"]): hdr[i].text = h
        for _, r in audit_df.iterrows():
            cells = table.add_row().cells
            vals = [r.get("item", ""), r.get("estado", ""), r.get("severidad", ""), r.get("observacion", ""), r.get("recomendacion", "")]
            for i, v in enumerate(vals): cells[i].text = str(v)
    if adoption_df is not None and not adoption_df.empty:
        doc.add_heading("Caudal adoptado", level=1)
        cols = list(adoption_df.columns)[:10]
        table = doc.add_table(rows=1, cols=len(cols))
        for i, c in enumerate(cols): table.rows[0].cells[i].text = str(c)
        for _, r in adoption_df.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(cols): cells[i].text = str(r.get(c, ""))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def pdf_report_bytes(title: str, audit_df: pd.DataFrame, score_df: pd.DataFrame, adoption_df: pd.DataFrame | None = None) -> bytes:
    """Creates a compact PDF report if reportlab is installed."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles['Title']), Spacer(1, 12)]
    if score_df is not None and not score_df.empty:
        row = score_df[score_df.get("tipo", "") == "resumen"].head(1)
        if row.empty: row = score_df.head(1)
        r = row.iloc[0]
        story += [Paragraph(f"Nota global: <b>{r.get('nota_global', 'NA')} / 10</b>", styles['Normal']), Paragraph(f"Estado final: <b>{r.get('estado_final', 'NA')}</b>", styles['Normal']), Spacer(1, 12)]
    if audit_df is not None and not audit_df.empty:
        story.append(Paragraph("Observaciones de auditoría", styles['Heading1']))
        data = [["Item", "Estado", "Severidad", "Observación"]]
        for _, r in audit_df.head(25).iterrows():
            data.append([str(r.get("item", ""))[:24], str(r.get("estado", ""))[:12], str(r.get("severidad", ""))[:10], str(r.get("observacion", ""))[:70]])
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.lightgrey),('GRID',(0,0),(-1,-1),0.25,colors.grey),('FONTSIZE',(0,0),(-1,-1),7)]))
        story.append(t)
    doc.build(story)
    return bio.getvalue()
